import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import sqlite3
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import numpy as np
import math
import pandas as pd
from docx import Document
from docx.shared import Inches
import tempfile

plt.rcParams['font.sans-serif']=['SimHei']

class SystemCalculationWindow:
    def __init__(self, master):
        self.master = master
        self.window = tk.Toplevel(master)
        self.window.title("系统计算及风机选型")
        self.window.geometry("800x800")
        
        # 连接项目数据库
        self.project_conn = sqlite3.connect('projects.db')
        self.project_cursor = self.project_conn.cursor()

        self.notebook = ttk.Notebook(self.window)
        self.notebook.pack(fill="both", expand=True)

        self.conn = sqlite3.connect('fan_models.db')
        self.cursor = self.conn.cursor()
        self.cursor.execute('CREATE TABLE IF NOT EXISTS fan_models (model_name TEXT, rated_pressure REAL, rated_flow REAL, motor_power REAL, motor_speed REAL)')

        self.fan_selection_tab = tk.Frame(self.notebook)
        self.notebook.add(self.fan_selection_tab, text="风机选型")
        self.create_fan_selection_widgets()

        self.data_management_tab = tk.Frame(self.notebook)
        self.notebook.add(self.data_management_tab, text="风机数据管理")
        self.create_data_management_widgets()

        self.plot_tab = tk.Frame(self.notebook)
        self.notebook.add(self.plot_tab, text="绘图")
        self.create_plot_widgets()
        
        self.selected_project = None
        self.matched_fan_models_for_plot = []
        
    def get_projects(self):
        self.project_cursor.execute('SELECT company_name, contact_name, contact_phone, project_name, project_location FROM projects')
        projects = self.project_cursor.fetchall()
        project_list = []
        for project in projects:
            project_dict = {
                "company_name": project[0],
                "contact_name": project[1],
                "contact_phone": project[2],
                "project_name": project[3],
                "project_location": project[4]
            }
            project_list.append(project_dict)
        return project_list
        
    def select_project(self):
        # 创建一个弹出窗口让用户选择项目
        project_window = tk.Toplevel(self.window)
        project_window.title("选择项目")

        # 创建一个下拉菜单
        project_label = ttk.Label(project_window, text="选择项目：")
        project_label.pack(pady=5)

        project_combobox = ttk.Combobox(project_window, values=[p["project_name"] for p in self.get_projects()])
        project_combobox.pack(pady=5)
        project_combobox.set("请选择项目")  # 设置默认值

        # 确定按钮
        def confirm_selection():
            selected_project_name = project_combobox.get()
            for project in self.get_projects():
                if project["project_name"] == selected_project_name:
                    self.selected_project = project
                    project_window.destroy()
                    return project
            messagebox.showerror("错误", "请选择一个有效的项目")

        confirm_button = ttk.Button(project_window, text="确定", command=confirm_selection)
        confirm_button.pack(pady=5)

        # 防止用户关闭窗口而不选择项目
        project_window.grab_set()
        project_window.wait_window()

        # 返回选中的项目信息
        return self.selected_project
    def calculate_energy_consumption(self, motor_power):
    # 假设运行时间为1小时
        running_time = 1
        return motor_power * running_time

    def create_fan_selection_widgets(self):
        self.diameter_label = tk.Label(self.fan_selection_tab, text="风管直径(mm)：")
        self.diameter_label.pack()
        self.diameter_entry = tk.Entry(self.fan_selection_tab)
        self.diameter_entry.pack()

        self.altitude_label = tk.Label(self.fan_selection_tab, text="海拔高度(m)：")
        self.altitude_label.pack()
        self.altitude_entry = tk.Entry(self.fan_selection_tab)
        self.altitude_entry.pack()

        self.slope_label = tk.Label(self.fan_selection_tab, text="斜井坡度(%)：")
        self.slope_label.pack()
        self.slope_entry = tk.Entry(self.fan_selection_tab)
        self.slope_entry.pack()

        self.temperature_label = tk.Label(self.fan_selection_tab, text="环境温度(℃)：")
        self.temperature_label.pack()
        self.temperature_entry = tk.Entry(self.fan_selection_tab)
        self.temperature_entry.pack()

        self.wind_speed_label = tk.Label(self.fan_selection_tab, text="风速(m/s)：")
        self.wind_speed_label.pack()
        self.wind_speed_entry = tk.Entry(self.fan_selection_tab)
        self.wind_speed_entry.pack()

        self.calculate_button = tk.Button(self.fan_selection_tab, text="开始计算", command=self.calculate_fan_selection)
        self.calculate_button.pack()
        
        self.export_fan_selection_button = tk.Button(self.fan_selection_tab, text="导出风机选型性能曲线到Word",
                                                     command=self.export_fan_selection_curve_to_word)
        self.export_fan_selection_button.pack()

        self.result_text = tk.Text(self.fan_selection_tab, height=8, width=50)
        self.result_text.pack()

        self.fig, self.ax = plt.subplots()
        self.canvas = FigureCanvasTkAgg(self.fig, master=self.fan_selection_tab)
        self.canvas.get_tk_widget().pack()
        

    def calculate_fan_selection(self):
        try:
            diameter_mm = float(self.diameter_entry.get())
            elevation_m = float(self.altitude_entry.get())
            slope_percent = float(self.slope_entry.get())
            temperature_celsius = float(self.temperature_entry.get())
            wind_speed_mps = float(self.wind_speed_entry.get())

            flow_rate, pressure = self.calculate_pressure_and_flow(diameter_mm, elevation_m, slope_percent, temperature_celsius, wind_speed_mps)

            # 四舍五入为整数
            flow_rate = int(round(flow_rate))
            pressure = int(round(pressure))

            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, f"计算得出的风量：{flow_rate} m³/s\n计算得出的风压：{pressure} Pa\n")

            self.cursor.execute('SELECT * FROM fan_models')
            fan_models = self.cursor.fetchall()

            suitable_models = []
            for model in fan_models:
                model_name, rated_pressure, rated_flow, motor_power, motor_speed = model
                # 将数据库中的额定风压和额定风量转换为整数
                rated_pressure = int(round(rated_pressure))
                rated_flow = int(round(rated_flow))
                if rated_pressure >= pressure and rated_flow >= flow_rate * 3600:
                    suitable_models.append((model_name, rated_pressure, rated_flow, motor_power, motor_speed))

            if suitable_models:
                self.ax.clear()
                for model in suitable_models:
                    model_name, rated_pressure, rated_flow, motor_power, motor_speed = model
                    # 根据模型数据生成性能曲线数据点（这里简单模拟，实际需根据具体模型数据计算）
                    flow_values = np.linspace(rated_flow * 0.5, rated_flow * 1.5, 100)
                    pressure_values = rated_pressure * (1 - (flow_values - rated_flow) / rated_flow) ** 2
                    self.ax.plot(flow_values, pressure_values, label=f'{model_name}')

                self.ax.set_xlabel('风量(m³/h)')
                self.ax.set_ylabel('风压(Pa)')
                self.ax.set_title('风机性能曲线')
                self.ax.legend()
                self.canvas.draw()

                selected_model = suitable_models[0]
                model_name, rated_pressure, rated_flow, motor_power, motor_speed = selected_model
                energy_consumption = self.calculate_energy_consumption(motor_power)
                result_text = f"推荐风机型号：{model_name}\n额定风压：{rated_pressure} Pa\n额定风量：{rated_flow} m³/h\n电机额定功率：{motor_power} kW\n电机转速：{motor_speed} rpm\n每小时能耗：{energy_consumption} kWh"
                self.result_text.insert(tk.END, result_text)
            else:
                self.result_text.insert(tk.END, "未找到合适的风机型号。")
        except ValueError:
            messagebox.showerror("输入错误", "请输入有效的数值")

    def calculate_pressure_and_flow(self, diameter_mm, elevation_m, slope_percent, temperature_celsius, wind_speed_mps):
        # 将直径从毫米转换为米
        diameter_m = diameter_mm / 1000.0
        # 计算风管截面积
        area_m2 = math.pi * (diameter_m ** 2) / 4
        # 计算风量
        flow_rate_m3s = area_m2 * wind_speed_mps

        # 计算空气密度
        rho_0 = 1.225  # 海平面上的空气密度
        T = temperature_celsius + 273.15  # 环境温度转换为开尔文
        rho = rho_0 * (1 - 0.0065 * elevation_m / (T + 0.0065 * elevation_m + 273.15)) ** 1.225

        # 计算风压
        g = 9.81  # 重力加速度
        pressure_pascal = rho * g * elevation_m + 0.5 * rho * (wind_speed_mps ** 2)

        return flow_rate_m3s, pressure_pascal

    def create_data_management_widgets(self):
        self.data_frame = tk.LabelFrame(self.data_management_tab, text="风机数据管理")
        self.data_frame.pack(fill="both", padx=10, pady=10)

        self.tree = ttk.Treeview(self.data_frame, columns=("型号", "额定风压", "额定风量", "功率", "转速"),
                                 show="headings")
        self.tree.pack(side="left", fill="both", expand=True)
        for col in self.tree['columns']:
            self.tree.heading(col, text=col)
        self.load_data()

        self.scrollbar = ttk.Scrollbar(self.data_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscroll=self.scrollbar.set)
        self.scrollbar.pack(side="right", fill="y")

        self.button_frame = tk.Frame(self.data_management_tab)
        self.button_frame.pack(fill="x", padx=10, pady=5)

        self.add_button = tk.Button(self.button_frame, text="添加", command=self.add_fan)
        self.add_button.pack(side="left", padx=5)

        self.delete_button = tk.Button(self.button_frame, text="删除", command=self.delete_fan)
        self.delete_button.pack(side="left", padx=5)

        self.import_button = tk.Button(self.button_frame, text="导入", command=self.import_data)
        self.import_button.pack(side="left", padx=5)

        self.export_button = tk.Button(self.button_frame, text="导出", command=self.export_data)
        self.export_button.pack(side="left", padx=5)

    def load_data(self):
        for row in self.tree.get_children():
            self.tree.delete(row)
        self.cursor.execute('SELECT * FROM fan_models')
        for row in self.cursor.fetchall():
            self.tree.insert("", "end", values=row)

    def add_fan(self):
        def save_fan():
            model_name = entry_model_name.get()
            rated_pressure = entry_pressure.get()
            rated_flow = entry_flow.get()
            motor_power = entry_power.get()
            motor_speed = entry_speed.get()

            if model_name and rated_pressure and rated_flow and motor_power and motor_speed:
                self.cursor.execute('INSERT INTO fan_models VALUES (?,?,?,?,?)',
                                    (model_name, float(rated_pressure), float(rated_flow), float(motor_power),
                                     float(motor_speed)))
                self.conn.commit()
                self.load_data()
                add_window.destroy()
            else:
                messagebox.showwarning("输入错误", "请填写所有字段")

        add_window = tk.Toplevel(self.master)
        add_window.title("添加风机")

        tk.Label(add_window, text="型号:").grid(row=0, column=0, padx=10, pady=5)
        entry_model_name = tk.Entry(add_window)
        entry_model_name.grid(row=0, column=1, padx=10, pady=5)

        tk.Label(add_window, text="额定风压 (Pa):").grid(row=1, column=0, padx=10, pady=5)
        entry_pressure = tk.Entry(add_window)
        entry_pressure.grid(row=1, column=1, padx=10, pady=5)

        tk.Label(add_window, text="额定风量 (m³/h):").grid(row=2, column=0, padx=10, pady=5)
        entry_flow = tk.Entry(add_window)
        entry_flow.grid(row=2, column=1, padx=10, pady=5)

        tk.Label(add_window, text="功率 (kW):").grid(row=3, column=0, padx=10, pady=5)
        entry_power = tk.Entry(add_window)
        entry_power.grid(row=3, column=1, padx=10, pady=5)
        tk.Label(add_window, text="转速 (rpm):").grid(row=4, column=0, padx=10, pady=5)
        entry_speed = tk.Entry(add_window)
        entry_speed.grid(row=4, column=1, padx=10, pady=5)

        tk.Button(add_window, text="保存", command=save_fan).grid(row=5, column=0, columnspan=2, pady=10)

    def delete_fan(self):
        selected_item = self.tree.selection()
        if selected_item:
            for item in selected_item:
                model_name = self.tree.item(item, 'values')[0]
                self.cursor.execute('DELETE FROM fan_models WHERE model_name=?', (model_name,))
                self.conn.commit()
                self.tree.delete(item)
        else:
            messagebox.showwarning("删除错误", "请选择要删除的风机型号")

    def import_data(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel 文件", "*.xlsx"), ("CSV 文件", "*.csv")])
        if file_path:
            try:
                if file_path.endswith('.xlsx'):
                    data = pd.read_excel(file_path)
                else:
                    data = pd.read_csv(file_path)

                for _, row in data.iterrows():
                    self.cursor.execute('INSERT INTO fan_models VALUES (?,?,?,?,?)',
                                        (row['型号'], row['额定风压'], row['额定风量'], row['功率'], row['转速']))
                self.conn.commit()
                self.load_data()
                messagebox.showinfo("导入成功", "数据已成功导入")
            except Exception as e:
                messagebox.showerror("导入错误", str(e))

    def export_data(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".xlsx",
                                                 filetypes=[("Excel 文件", "*.xlsx"), ("CSV 文件", "*.csv")])
        if file_path:
            try:
                self.cursor.execute('SELECT * FROM fan_models')
                data = pd.DataFrame(self.cursor.fetchall(), columns=['型号', '额定风压', '额定风量', '功率', '转速'])

                if file_path.endswith('.xlsx'):
                    data.to_excel(file_path, index=False)
                else:
                    data.to_csv(file_path, index=False)

                messagebox.showinfo("导出成功", "数据已成功导出")
            except Exception as e:
                messagebox.showerror("导出错误", str(e))
                
    def export_fan_selection_curve_to_word(self):
        if self.ax.lines:
            doc = Document()
            doc.add_heading('风机选型报告', level=1)

            # 风机型号、性能参数
            doc.add_heading('风机型号及性能参数', level=2)
            if self.result_text.get("1.0", tk.END).strip():
                result = self.result_text.get("1.0", tk.END).strip().split('\n')
                for line in result:
                    doc.add_paragraph(line)

            # 风机性能曲线
            doc.add_heading('风机性能曲线', level=2)
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmpfile:
                self.fig.savefig(tmpfile.name)
            doc.add_picture(tmpfile.name, width=Inches(6))

            # 隧道施工参数
            doc.add_heading('隧道施工参数', level=2)
            diameter_mm = self.diameter_entry.get()
            elevation_m = self.altitude_entry.get()
            slope_percent = self.slope_entry.get()
            temperature_celsius = self.temperature_entry.get()
            wind_speed_mps = self.wind_speed_entry.get()
            doc.add_paragraph(f"风管直径: {diameter_mm} mm")
            doc.add_paragraph(f"海拔高度: {elevation_m} m")
            doc.add_paragraph(f"斜井坡度: {slope_percent} %")
            doc.add_paragraph(f"环境温度: {temperature_celsius} ℃")
            doc.add_paragraph(f"风速: {wind_speed_mps} m/s")

            # 项目信息
            doc.add_heading('项目信息', level=2)
            selected_project = self.select_project()
            if selected_project:
                doc.add_paragraph(f"甲方公司: {selected_project['company_name']}")
                doc.add_paragraph(f"甲方联系人: {selected_project['contact_name']}")
                doc.add_paragraph(f"甲方联系人电话: {selected_project['contact_phone']}")
                doc.add_paragraph(f"项目名称: {selected_project['project_name']}")
                doc.add_paragraph(f"项目地点: {selected_project['project_location']}")

            file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                     filetypes=[("Word 文件", "*.docx")])
            if file_path:
                doc.save(file_path)
                messagebox.showinfo("导出成功", "风机选型报告已成功导出到Word文档。")
            else:
                messagebox.showwarning("导出失败", "未指定保存路径。")

    def export_plot_curve_to_word(self):
        if self.matched_fan_models_for_plot:
            doc = Document()
            doc.add_heading('绘图性能曲线及风机模型参数', level=1)

            # 风机型号、性能参数
            doc.add_heading('匹配的风机模型参数', level=2)
            for model in self.matched_fan_models_for_plot:
                model_name, rated_pressure, rated_flow, motor_power, motor_speed = model
                doc.add_paragraph(f"型号: {model_name}")
                doc.add_paragraph(f"额定风压: {rated_pressure} Pa")
                doc.add_paragraph(f"额定风量: {rated_flow} m³/h")
                doc.add_paragraph(f"电机额定功率: {motor_power} kW")
                doc.add_paragraph(f"电机转速: {motor_speed} rpm")
                doc.add_paragraph("-" * 50)

        # 绘图性能曲线
            doc.add_heading('绘图性能曲线', level=2)
            fig, ax = plt.subplots()
            for model in self.matched_fan_models_for_plot:
                model_name, rated_pressure, rated_flow, motor_power, motor_speed = model
                flow_values = np.linspace(rated_flow * 0.5, rated_flow * 1.5, 100)
                pressure_values = rated_pressure * (1 - (flow_values - rated_flow) / rated_flow) ** 2
                ax.plot(flow_values, pressure_values, label=f'{model_name}')

            ax.set_xlabel('风量(m³/h)')
            ax.set_ylabel('风压(Pa)')
            ax.set_title('绘图性能曲线')
            ax.legend()

            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmpfile:
                fig.savefig(tmpfile.name)
            doc.add_picture(tmpfile.name, width=Inches(6))

            file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                 filetypes=[("Word 文件", "*.docx")])
            if file_path:
                doc.save(file_path)
                messagebox.showinfo("导出成功", "绘图性能曲线及风机模型参数已成功导出到Word文档。")
            else:
                messagebox.showwarning("导出失败", "未指定保存路径。")
        else:
            messagebox.showinfo("提示", "没有找到匹配的风机型号，无法导出绘图性能曲线。")

    def create_plot_widgets(self):
        self.plot_frame = tk.LabelFrame(self.plot_tab, text="风机性能曲线")
        self.plot_frame.pack(fill="both", padx=10, pady=10)

        self.export_plot_button = tk.Button(self.plot_tab, text="导出绘图性能曲线到Word",
                                            command=self.export_plot_curve_to_word)
        self.export_plot_button.pack()

        tk.Label(self.plot_frame, text="额定风压 (Pa):").grid(row=0, column=0, padx=10, pady=5)
        self.entry_pressure = tk.Entry(self.plot_frame)
        self.entry_pressure.grid(row=0, column=1, padx=10, pady=5)

        tk.Label(self.plot_frame, text="额定风量 (m³/h):").grid(row=1, column=0, padx=10, pady=5)
        self.entry_flow = tk.Entry(self.plot_frame)
        self.entry_flow.grid(row=1, column=1, padx=10, pady=5)

        self.plot_button = tk.Button(self.plot_frame, text="绘制性能曲线", command=self.plot_performance_curve)
        self.plot_button.grid(row=2, column=0, columnspan=2, pady=10)

        self.plot_canvas_frame = tk.Frame(self.plot_frame)
        self.plot_canvas_frame.grid(row=3, column=0, columnspan=2)

    def plot_performance_curve(self):
        try:
            pressure = float(self.entry_pressure.get())
            flow = float(self.entry_flow.get())

            self.cursor.execute('SELECT * FROM fan_models')
            fan_models = self.cursor.fetchall()

            self.matched_fan_models_for_plot = []
            for model in fan_models:
                model_name, rated_pressure, rated_flow, motor_power, motor_speed = model
                # 将数据库中的额定风压和额定风量转换为整数
                rated_pressure = int(round(rated_pressure))
                rated_flow = int(round(rated_flow))
                if rated_pressure >= pressure and rated_flow >= flow:
                    self.matched_fan_models_for_plot.append((model_name, rated_pressure, rated_flow, motor_power, motor_speed))

            if self.matched_fan_models_for_plot:
                fig, ax = plt.subplots()
                for model in self.matched_fan_models_for_plot:
                    model_name, rated_pressure, rated_flow, motor_power, motor_speed = model
                    # 根据模型数据生成性能曲线数据点（这里简单模拟，实际需根据具体模型数据计算）
                    flow_values = np.linspace(rated_flow * 0.5, rated_flow * 1.5, 100)
                    pressure_values = rated_pressure * (1 - (flow_values - rated_flow) / rated_flow) ** 2
                    ax.plot(flow_values, pressure_values, label=f'{model_name}')

                ax.set_xlabel('风量(m³/h)')
                ax.set_ylabel('风压(Pa)')
                ax.set_title('风机性能曲线')
                ax.legend()

                for widget in self.plot_canvas_frame.winfo_children():
                    widget.destroy()

                canvas = FigureCanvasTkAgg(fig, master=self.plot_canvas_frame)
                canvas.get_tk_widget().pack()
                canvas.draw()
            else:
                messagebox.showinfo("提示", "没有找到匹配的风机型号")
        except ValueError:
            messagebox.showerror("输入错误", "请正确填写额定风压和风量")


if __name__ == "__main__":
    root = tk.Tk()
    app = SystemCalculationWindow(root)
    root.mainloop